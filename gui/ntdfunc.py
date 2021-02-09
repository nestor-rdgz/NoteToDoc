'''
    Name: ntdfunc.py
    GitHub URL: https://github.com/nestor-rdgz/NoteToDoc
    License: GNU General Public License v3.0
'''
from pandas import read_csv
from bs4 import BeautifulSoup
from docx import Document      #Documentation: https://python-docx.readthedocs.io/en/latest/

def Notetodoc(address, docname, csvfile):

    biblio = read_csv(csvfile)
    
    if docname != "notes.docx":
        if docname.find(".docx")==(-1):
            docname = docname+".docx"
    else:
        docname = "notes.docx"


    biblio_sim = biblio[['Item Type', 'Publication Year', 'Title', 'Author', 'Notes', 'Url']]
    biblio_sim = biblio_sim.fillna('empty')

    shape = biblio_sim.shape
    rows = shape[0]

    #The .docx file is initiated
    documento = Document()
    documento.add_heading('Notes from references', 0)

    for i in range(rows):
        #Each row is a bibliography element
        document = biblio_sim.iloc[i] 
        item_type = document[0]
        year = str(document[1])
        title = document [2]
        authors = document[3]
        notes = document[4]
        url = document[5]

        #As the notes are in HTML format, it is necessary to extract only the text
        soup = BeautifulSoup(notes, features='lxml')
        notes_text = soup.get_text('\n').replace('\n \n', '\n').replace('\n \n \n', '\n')
            
        #The information is written into the .docx file
        #The title is inserted
        documento.add_heading("Title: "+title, 1)

        #The Author is inserted
        p = documento.add_paragraph("")
        p.add_run("Authors: ").bold = True
        p.add_run(authors).italic = True

        #The Year, publication type and url are inserted
        p = documento.add_paragraph("")
        p.add_run("Year: ").bold = True
        p.add_run(year[0:4])
        p.add_run("\nPublication type: ").bold = True
        p.add_run(item_type)
        p.add_run("\nURL: ").bold = True
        p.add_run(url)

        #The Notes are inserted
        p = documento.add_paragraph("")
        p.add_run("NOTES: \n").bold = True
        p.add_run(notes_text)
        
    #Closing file
    documento.save(address+"/"+docname)
    

